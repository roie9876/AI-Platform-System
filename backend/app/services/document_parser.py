from __future__ import annotations

import hashlib
import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)


class TextChunker:
    """Split text into overlapping chunks for RAG indexing."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk(self, text: str) -> List[str]:
        """Split text into chunks with overlap. Uses paragraph splitting first,
        then sentence splitting if paragraphs are too large."""
        if not text.strip():
            return []

        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        chunks: List[str] = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk) + len(para) + 2 <= self.chunk_size:
                current_chunk = f"{current_chunk}\n\n{para}" if current_chunk else para
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                if len(para) > self.chunk_size:
                    sentence_chunks = self._split_by_sentences(para)
                    chunks.extend(sentence_chunks)
                    current_chunk = ""
                else:
                    current_chunk = para

        if current_chunk:
            chunks.append(current_chunk)

        if self.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks)

        return chunks

    def _split_by_sentences(self, text: str) -> List[str]:
        """Split text by sentences when paragraph is too large."""
        sentences = text.replace(". ", ".\n").split("\n")
        chunks: List[str] = []
        current = ""
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            if len(current) + len(sentence) + 1 <= self.chunk_size:
                current = f"{current} {sentence}" if current else sentence
            else:
                if current:
                    chunks.append(current)
                current = sentence
        if current:
            chunks.append(current)
        return chunks

    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """Add overlapping text between consecutive chunks."""
        result = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_text = chunks[i - 1]
            overlap = (
                prev_text[-self.chunk_overlap :]
                if len(prev_text) > self.chunk_overlap
                else prev_text
            )
            result.append(f"{overlap} {chunks[i]}")
        return result


class DocumentParser:
    """Parse documents of various formats into plain text."""

    SUPPORTED_TYPES = {"pdf", "txt", "md", "docx"}

    def parse(self, file_path: str, content_type: str | None = None) -> str:
        """Parse a document file and return extracted text."""
        path = Path(file_path)
        ext = path.suffix.lower().lstrip(".")

        if ext == "pdf":
            return self._parse_pdf(path)
        elif ext in ("txt", "md"):
            return self._parse_text(path)
        elif ext == "docx":
            return self._parse_docx(path)
        else:
            raise ValueError(
                f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_TYPES}"
            )

    def parse_bytes(self, content: bytes, filename: str) -> str:
        """Parse document from bytes (for upload handling)."""
        ext = Path(filename).suffix.lower().lstrip(".")

        if ext == "pdf":
            return self._parse_pdf_bytes(content)
        elif ext in ("txt", "md"):
            return content.decode("utf-8", errors="replace")
        elif ext == "docx":
            return self._parse_docx_bytes(content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def compute_hash(self, content: bytes) -> str:
        """Compute SHA-256 hash of file content for deduplication."""
        return hashlib.sha256(content).hexdigest()

    def _parse_pdf(self, path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    def _parse_pdf_bytes(self, content: bytes) -> str:
        import io

        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)

    def _parse_text(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    def _parse_docx(self, path: Path) -> str:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        return "\n\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )

    def _parse_docx_bytes(self, content: bytes) -> str:
        import io

        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        return "\n\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )
